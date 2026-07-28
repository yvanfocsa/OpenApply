#!/usr/bin/env python3
"""Validate a generated DOCX against its source, render it, and enforce one page."""

from __future__ import annotations

import argparse
import hashlib
import json
import shutil
import subprocess
import sys
import zipfile
from pathlib import Path

from docx import Document
from pypdf import PdfReader


def section_signatures(document: Document) -> list[tuple[str, ...]]:
    return [
        tuple(str(value) for value in (
            section.page_width,
            section.page_height,
            section.top_margin,
            section.bottom_margin,
            section.left_margin,
            section.right_margin,
            section.header_distance,
            section.footer_distance,
            section.orientation,
        ))
        for section in document.sections
    ]


def media_hashes(docx_path: Path) -> dict[str, str]:
    with zipfile.ZipFile(docx_path) as archive:
        return {
            name: hashlib.sha256(archive.read(name)).hexdigest()
            for name in archive.namelist()
            if name.startswith("word/media/")
        }


FORBIDDEN_FRENCH_IN_ENGLISH = [
    "alternance",
    "candidature",
    "objet :",
    "madame,",
    "monsieur,",
    "salutations distinguées",
]

FORBIDDEN_ENGLISH_IN_FRENCH = [
    "dear hiring manager",
    "permanent position",
]


def check_language_purity(doc: Document, language: str = "auto") -> None:
    full_text = "\n".join(p.text for p in doc.paragraphs).lower()
    is_english = language == "en" or ("profile" in full_text and "professional experience" in full_text)
    is_french = language == "fr" or ("profil" in full_text and "expériences professionnelles" in full_text)
    if is_english:
        for forbidden in FORBIDDEN_FRENCH_IN_ENGLISH:
            if forbidden in full_text:
                raise RuntimeError(f"Language purity check failed: French term '{forbidden}' found in English document")
    elif is_french:
        for forbidden in FORBIDDEN_ENGLISH_IN_FRENCH:
            if forbidden in full_text:
                raise RuntimeError(f"Language purity check failed: English term '{forbidden}' found in French document")


def validate_structure(template_path: Path, docx_path: Path, language: str = "auto") -> None:
    if template_path.resolve() == docx_path.resolve():
        raise RuntimeError("The generated document must not overwrite its source template")
    source = Document(template_path)
    generated = Document(docx_path)
    if len(generated.sections) != len(source.sections):
        raise RuntimeError("Document section count changed")
    if section_signatures(generated) != section_signatures(source):
        raise RuntimeError("Document page geometry changed")
    if len(generated.paragraphs) != len(source.paragraphs):
        raise RuntimeError("Top-level paragraph count changed")
    if len(generated.tables) != len(source.tables):
        raise RuntimeError("Table count changed")
    if len(generated._element.xpath(".//w:drawing")) != len(source._element.xpath(".//w:drawing")):
        raise RuntimeError("Embedded photo or drawing element was removed from document structure")
    if media_hashes(docx_path) != media_hashes(template_path):
        raise RuntimeError("Embedded images or media changed")
    check_language_purity(generated, language)


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--template", type=Path, required=True)
    parser.add_argument("--docx", type=Path, required=True)
    parser.add_argument("--pdf", type=Path, required=True)
    parser.add_argument("--qa-dir", type=Path, required=True)
    parser.add_argument("--label", default="document")
    parser.add_argument("--language", default="auto")
    args = parser.parse_args()

    template_path = args.template.expanduser().resolve()
    docx_path = args.docx.expanduser().resolve()
    if not template_path.is_file() or not docx_path.is_file():
        raise SystemExit(f"Missing source or generated DOCX for {args.label}")
    validate_structure(template_path, docx_path, args.language)

    qa_dir = args.qa_dir.expanduser().resolve()
    qa_dir.mkdir(parents=True, exist_ok=True)
    for candidate in qa_dir.glob("page-*.png"):
        candidate.unlink()
    emitted_pdf = qa_dir / f"{docx_path.stem}.pdf"
    emitted_pdf.unlink(missing_ok=True)
    renderer = Path(__file__).with_name("render_docx.py")
    completed = subprocess.run([
        sys.executable,
        str(renderer),
        str(docx_path),
        "--output_dir",
        str(qa_dir),
        "--emit_pdf",
    ], text=True, capture_output=True, check=False)
    if completed.returncode:
        raise SystemExit(completed.stderr or completed.stdout or f"Rendering failed for {args.label}")
    if not emitted_pdf.is_file():
        raise SystemExit(f"Renderer did not emit a PDF for {args.label}")
    pages = len(PdfReader(emitted_pdf).pages)
    images = sorted(qa_dir.glob("page-*.png"))
    if pages != 1 or len(images) != 1:
        raise SystemExit(f"ONE-PAGE GATE FAILED for {args.label}: PDF={pages}, images={len(images)}")
    args.pdf.expanduser().resolve().parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(emitted_pdf, args.pdf.expanduser().resolve())
    print(json.dumps({
        "status": "ok",
        "label": args.label,
        "pdf": str(args.pdf.expanduser().resolve()),
        "pageImage": str(images[0]),
    }, ensure_ascii=False))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())